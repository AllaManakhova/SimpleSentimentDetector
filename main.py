import stanza
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from htmldocx import HtmlToDocx

from rusentilex_2017_handler import rusentilex_2017_handler


def get_text_file(file_name):
    return open(file_name, encoding='utf8').read()


def sentiment_analyze(text):
    rusentilex_2017 = rusentilex_2017_handler()
    html_text = text
    doc = nlp(text)
    stat = set()
    stat_dict = {}
    for sentence in doc.sentences:
        for word in sentence.words:
            for d in rusentilex_2017:
                if d["word"] == word.lemma:
                    if d[" sentiment"] == " negative":
                        positive_str = f'<span style=\" font-size:8pt; font-weight:600; color:#ff0000;\" >{word.text}</span>'
                        html_text = html_text.replace(word.text, positive_str, 1)
                        stat.add(
                            str(word.text) + " - " + str(d[" part_of_speech"]).lower() + " - " + str(d[" sentiment"]))
                    elif d[" sentiment"] == " positive":
                        negative_str = f'<span style=\" font-size:8pt; font-weight:600; color:#00ff00;\" >{word.text}</span>'
                        html_text = html_text.replace(word.text, negative_str, 1)
                        stat.add(
                            str(word.text) + " - " + str(d[" part_of_speech"]).lower() + " - " + str(d[" sentiment"]))

    for elem in stat:
        elem = elem.split(" - ")[1]
        if elem in stat_dict.keys():
            stat_dict[elem] += 1
        else:
            stat_dict[elem] = 1
    return html_text, stat, stat_dict


def dict_to_str(stat_dict):
    res_str = ""
    for key, value in stat_dict.items():
        res_str = res_str + str(key) + " - " + str(value) + "\n"
    return res_str


def create_file(file_name, text, stat, stat_dict):
    document = Document()
    parser = HtmlToDocx()
    parser.add_html_to_document(text, document)
    document.add_paragraph().add_run("\n\nСТАТИСТИКА:\n").underline = True
    if len(stat) == 0:
        document.add_paragraph("Тонально окрашеных слов нет")
    else:
        for elem in stat:
            if elem.split(" - ")[2] == " negative":
                document.add_paragraph().add_run(elem).font.highlight_color = WD_COLOR_INDEX.RED
            elif elem.split(" - ")[2] == " positive":
                document.add_paragraph().add_run(elem).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        document.add_paragraph(stat_dict)
    document.save(file_name + ".docx")


def main():
    list_of_files = []
    for file in list_of_files:
        text = get_text_file('data/' + file)
        html_text, stat, stat_dict = sentiment_analyze(text)
        res_name = 'res/res-' + file.split('.')[0]
        stat_dict_str = dict_to_str(stat_dict)
        create_file(res_name, html_text, stat, stat_dict_str)


stanza.download('ru', model_dir='resources/', processors='tokenize,ner,lemma,pos')
nlp = stanza.Pipeline(lang='ru', dir='resources/', processors='tokenize,ner,lemma,pos')
main()
